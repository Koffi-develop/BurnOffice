import os
from random import *
lc000 = os.getenv('TEMP')
loglocal = lc000 + "/PNLOG_" + str(randint(1111111, 9999999)) + ".tmp"
bb00 = open(loglocal, 'w')
bb00.write("BurnOffice Log \n")
bb00.close()
bb0 = open(loglocal, 'a')
with open(loglocal, 'a') as j:
    j.write("Loading Paths... \n")
with open(loglocal, 'a') as j:
    j.write("Loading Application (KDS engine)... \n")
from tkinter import *
import tkinter
from tkinter import ttk
import tkinter as tk
from tkinter.filedialog import *
from tkinter.messagebox import *
from random import *
import tkinter.font as tkFont
import pycbrf
import configparser
import csv
import subprocess
import docx2txt
from docx import Document
from odf import text, teletype
from odf.opendocument import load
from odf import opendocument
import pandas as pd
from pandastable import Table
from tkintertable import TableCanvas
global config
global gov000
global mainconfigsector
global fontsector
global defaultdocument
with open(loglocal, 'a') as j:
    j.write("Loading Locales... \n")
root = Tk() 
root.title("BurnOffice")
root.geometry("700x500")
gov000 = 0
config = configparser.ConfigParser()
config.read('mains.ini')
mainconfigsector = config['main']
mainconfigsector = config['font']
defaultdocument = config['main']['defaultdocument']
with open(loglocal, 'a') as j:
    j.write("Constructing logic... \n")
def myopen():
    gov000 = 0
    global qq
    qq = askopenfilename()
    if qq == '':
        print('False')
    else:
        filename, file_extension = os.path.splitext(qq)
        if file_extension == ".txt":
            titl3 = "BurnOffice - " + qq
            root.title(titl3)
            with open(qq, 'r') as f:
                text = f.read()
                txt.delete('1.0', END)
                txt.insert(END, text)
                config.set('main','defaultdocument', qq)
                with open('mains.ini', 'w') as configfile:
                    config.write(configfile)
                config.set('main','fileextension', '.txt')
                with open('mains.ini', 'w') as configfile:
                    config.write(configfile)
                with open(loglocal, 'a') as j:
                    j.write("File opened: " + qq + " \n")
        if file_extension == ".TXT":
            titl3 = "BurnOffice - " + qq
            root.title(titl3)
            with open(qq, 'r') as f:
                text = f.read()
                txt.delete('1.0', END)
                txt.insert(END, text)
                config.set('main','defaultdocument', qq)
                with open('mains.ini', 'w') as configfile:
                    config.write(configfile)
                config.set('main','fileextension', '.TXT')
                with open('mains.ini', 'w') as configfile:
                    config.write(configfile)
                with open(loglocal, 'a') as j:
                    j.write("File opened: " + qq + " \n")
        elif file_extension == ".csv":
            nfl5 = Toplevel()
            nfl5.geometry('800x600')
            try:
                file_name = qq
                df = pd.read_csv(file_name)        
                if (len(df)== 0):
                    msg.showinfo('No records', 'No records')
                else:
                    pass
                f2 = Frame(nfl5, height=200, width=300) 
                f2.pack(fill=BOTH,expand=1)
                table = Table(f2, dataframe=df,read_only=True)
                table.show()
                config.set('main','defaultdocument', qq)
                with open('mains.ini', 'w') as configfile:
                    config.write(configfile)
                config.set('main','fileextension', '.csv')
                with open('mains.ini', 'w') as configfile:
                    config.write(configfile)
            except FileNotFoundError as e:
                print(e)
                showerror('Error in opening file',e) 
        elif file_extension == ".xls":
            nfl6 = Toplevel()
            nfl6.geometry('800x600')
            try:
                file_name = qq
                df = pd.read_excel(file_name)        
                if (len(df)== 0):
                    msg.showinfo('No records', 'No records')
                else:
                    pass
                f2 = Frame(nfl6, height=200, width=300) 
                f2.pack(fill=BOTH,expand=1)
                table = Table(f2, dataframe=df,read_only=True)
                table.show()
                config.set('main','defaultdocument', qq)
                with open('mains.ini', 'w') as configfile:
                    config.write(configfile)
                config.set('main','fileextension', '.xls')
                with open('mains.ini', 'w') as configfile:
                    config.write(configfile)
            except FileNotFoundError as e:
                print(e)
                showerror('Error in opening file',e) 
        elif file_extension == ".docx":
            ee = docx2txt.process(qq)
            txt.delete('1.0', END)
            txt.insert(END, ee)
            config.set('main','defaultdocument', qq)
            with open('mains.ini', 'w') as configfile:
                config.write(configfile)
            config.set('main','fileextension', '.docx')
            with open('mains.ini', 'w') as configfile:
                    config.write(configfile)
        else:
            root.title('Incorrect File Extension!')
def mynew():
    gov000 = 0
    global qq
    up = os.getenv('USERPROFILE')
    qq = up +"/Desktop/PN_" + str(randint(1111111, 9999999)) + ".txt"
    root.title("BurnOffice - New File")
    txt.delete('1.0', END)
    with open(loglocal, 'a') as j:
        j.write("New File created: " + qq +" \n")
def mysaveas():
    gov000 = 1
    global sa
    sa = asksaveasfilename()
    if sa == '':
        print('False')
    else:
        if config['main']['fileextension'] == '.txt':
            with open(sa, 'w') as b:
                g = txt.get(1.0, END)
                b.write(g)
            with open(loglocal, 'a') as j:
                j.write("File saved as: " + sa + " \n")
        elif config['main']['fileextension'] == '.TXT':
            with open(sa, 'w') as b:
                g = txt.get(1.0, END)
                b.write(g)
            with open(loglocal, 'a') as j:
                j.write("File saved as: " + sa + " \n")
        elif config['main']['fileextension'] == '.docx':
            document = Document()
            def saveeeeeee():
                g = txt.get(1.0, END)
                document.add_paragraph(g)
                document.save(sa)
            if os.path.exists(qq) == True:
                os.remove(qq)
                saveeeeeee()
            else:
                saveeeeeee()
def mysave():
    gov000 = 1
    sa = qq
    if config['main']['fileextension'] == '.txt':
        with open(sa, 'w') as b:
            g = txt.get(1.0, END)
            b.write(g)
        with open(loglocal, 'a') as j:
            j.write("File saved as: " + sa + " \n")
    elif config['main']['fileextension'] == '.TXT':
        with open(sa, 'w') as b:
            g = txt.get(1.0, END)
            b.write(g)
        with open(loglocal, 'a') as j:
            j.write("File saved as: " + sa + " \n")
    elif config['main']['fileextension'] == '.docx':
        document = Document()
        def saveeeeeee():
            g = txt.get(1.0, END)
            document.add_paragraph(g)
            document.save(sa)
        if os.path.exists(qq) == True:
            os.remove(qq)
            saveeeeeee()
        else:
            saveeeeeee()  
def myexit():
    with open(qq, 'r') as i:
        global ioi
        global hend
        ioi = i.read()
        hend = txt.get('1.0', END)
    if ioi == hend:
        root.destroy()
    else:
        if askyesno("Exit",'Do you want to save?'):
            mysaveas()
            with open(loglocal, 'a') as j:
                j.write("Exit with save: " + qq + " \n")
            root.destroy()
        else:
            with open(loglocal, 'a') as j:
                j.write("Exit dont with save: " + qq + " \n")
            root.destroy()
def vl():
    with open(loglocal, 'a') as j:
        j.write('Loading Toplevel: nfl0 \n')
    nfl0 = Toplevel()
    nfl0.geometry("300x300")
    tex = Text(nfl0, width=2000, height=2000)
    tex.pack()
    tex.delete('1.0', END)
    with open(loglocal, 'r') as b0000:
        tex0000 = b0000.read()
        tex.insert(END, tex0000)
def fonti():
    nfl3 = Toplevel()
    nfl3.geometry("300x300")
    texs = Label(nfl3, text="Choose Font")
    choosen = ttk.Combobox(nfl3, width = 27)
    choosen['values'] = ('Arial', 
                          'Times New Roman',
                          'Comic Sans MS',
                          'Terminal',
                          'Courier',
                          'MV Boli', 
                          'Impact', 
                          'Tahoma', 
                          'Calibri', 
                          'System', 
                          'Cambria', 
                          'Corbel', 
                          'Consolas', 
                          'Verdana')
    texs2 = Label(nfl3, text="Choose Font Size")
    ent = Entry(nfl3, width=10)
    texs3 = Label(nfl3, text="Choose Thickness")
    choosen2 = ttk.Combobox(nfl3, width = 27)
    choosen2['values'] = ('Normal', 
                          'Bold')
    texs4 = Label(nfl3, text="Choose Incline")
    choosen3 = ttk.Combobox(nfl3, width = 27)
    choosen3['values'] = ('Italic', 
                          'Roman')
    buttt = Button(nfl3, text="OK")
    texs.grid()
    choosen.grid()
    texs3.grid()
    choosen2.grid()
    texs4.grid()
    choosen3.grid()
    texs2.grid()
    ent.grid()
    buttt.grid()
    if config['font']['font'] == 'Times':
        choosenfont = 'Times New Roman'
    else:
        choosenfont = config['font']['font']
    choosen.set(choosenfont)
    if config['font']['incline'] == 'normal':
        hdsjd = 'Normal'
    else:
        hdsjd = 'Bold'
    if config['font']['thickness'] == 'italic':
        uiwuw = 'Italic'
    else: 
        uiwuw = 'Roman'
    choosen2.set(hdsjd)
    choosen3.set(uiwuw)
    ent.delete(0,"end")
    ent.insert(0, config['font']['size'])
    def savefonti(event):
        choosenfont = choosen.get()
        choosenfont2 = choosen2.get()
        choosenfont3 = choosen3.get()
        if choosenfont == 'Times New Roman':
            choosenfont = 'Times'
        if choosenfont2 == 'Normal':
            choosenfont2 = 'normal'
        else:
            choosenfont2 = 'bold'
        if choosenfont3 == 'Italic':
            choosenfont3 = 'italic'
        else:
            choosenfont3 = 'roman'
        config.set('font','font', choosenfont)
        config.set('font','size', ent.get())
        config.set('font','thickness', choosenfont2)
        config.set('font','incline', choosenfont3)
        with open('mains.ini', 'w') as configfile:
            config.write(configfile)
        smile = tkFont.Font(family=choosenfont, size=ent.get(), weight=choosenfont2, slant=choosenfont3)
        txt.config(font = smile)
        with open(loglocal, 'a') as j:
            j.write('Font have been changed. \n')
        nfl3.destroy()
    buttt.bind('<Button-1>', savefonti)
def myrun():
    os.system(qq)
def colori():
    nfl4 = Toplevel()
    nfl4.geometry("300x300")
    texs = Label(nfl4, text="Choose Background color")
    choosen = ttk.Combobox(nfl4, width = 27)
    choosen['values'] = ('Green', 
                          'Lime',
                          'Yellow',
                          'Red',
                          'Blue',
                          'Cyan', 
                          'Purple', 
                          'Black', 
                          'Pink', 
                          'Orange', 
                          'White')
    texs2 = Label(nfl4, text="Choose Text color")
    choosen2 = ttk.Combobox(nfl4, width = 27)
    choosen2['values'] = ('Green', 
                          'Lime',
                          'Yellow',
                          'Red',
                          'Blue',
                          'Cyan', 
                          'Purple', 
                          'Black', 
                          'Pink', 
                          'Orange', 
                          'White')
    buttt = Button(nfl4, text="OK")
    texs.grid()
    choosen.grid()
    texs2.grid()
    choosen2.grid()
    buttt.grid()
    choosen.set(config['color']['bg'])
    choosen2.set(config['color']['fg'])
    def colorisave(event):
        choosencol = choosen.get()
        choosencol2 = choosen2.get()
        config.set('color','bg', choosencol)
        config.set('color','fg', choosencol2)
        with open('mains.ini', 'w') as configfile:
            config.write(configfile)
        txt.config(bg = config['color']['bg'])
        txt.config(fg = config['color']['fg'])
        with open(loglocal, 'a') as j:
            j.write('Color have been changed. \n')
        nfl4.destroy()
    buttt.bind('<Button-1>', colorisave)
def about():
    with open(loglocal, 'a') as j:
        j.write('Show Info: About \n')
    showinfo("BurnOffice", "Version 0.9 Alpha \n Copyright © KoffiDev \n All rights reserved.")
def absol():
    nfl2 = Toplevel()
    nfl2.geometry("100x75")
    ent55 = Entry(nfl2, width=50)
    ent55.pack()
    def op3n(event):
        global qq
        qq = ent55.get()
        if qq == '':
            print('False')
        else:
            titl3 = "BurnOffice - " + qq
            root.title(titl3)
            with open(qq, 'r') as f:
                text = f.read()
                txt.delete('1.0', END)
                txt.insert(END, text)
                with open(loglocal, 'a') as j:
                    j.write("File opened: " + qq + " \n")
                nfl2.destroy()
    but1 = Button(nfl2, text="Open")
    but1.pack()
    but1.bind("<Button-1>", op3n)
def non3():
    print('False')
def qtd():
    txt.delete('1.0', END)
    config.set('main','defaultdocument', '0')
    with open('mains.ini', 'w') as configfile:
        config.write(configfile)
    titl3 = "BurnOffice"
    root.title(titl3)
    with open(loglocal, 'a') as j:
        j.write("Quited from document... \n")
with open(loglocal, 'a') as j:
    j.write("Loading Objects... \n")
m = Menu(root)
root.config(menu=m)
fm = Menu(m)
m.add_cascade(label="File", menu = fm)
fm.add_command(label="New File", command = mynew)
fm.add_command(label="Open", command = myopen)
fm.add_command(label="Save As..", command = mysaveas)
fm.add_command(label="Save", command = mysave)
fm.add_command(label="Absolute path", command = absol)
fm.add_command(label="Quit this document", command = qtd)
fm.add_command(label="_________", command = non3)
fm.add_command(label="Exit", command = myexit)
hm = Menu(m)
am = Menu(m)
m.add_cascade(label="Options", menu=am)
am.add_command(label="Change Font", command = fonti)
am.add_command(label="Change Color", command = colori)
am.add_command(label="Run Current File", command = myrun)
am.add_command(label="View Log", command = vl)
m.add_cascade(label="Help", menu= hm)
hm.add_command(label="About", command = about)
txt = Text(root, width=2000, height=2000)
txt.pack()
scrollbar = Scrollbar(root)
scrollbar.pack(side=RIGHT, fill=Y)
txt.config(yscrollcommand=scrollbar.set)
scrollbar.config(command=txt.yview)
if defaultdocument == '0':
    print('False')
elif config['main']['fileextension'] == '.txt':
    with open(defaultdocument, 'r') as f:
        text = f.read()
        titl3 = "BurnOffice - " + defaultdocument
        root.title(titl3)
        txt.delete('1.0', END)
        txt.insert(END, text)
elif config['main']['fileextension'] == '.TXT':
    with open(defaultdocument, 'r') as f:
        text = f.read()
        titl3 = "BurnOffice - " + defaultdocument
        root.title(titl3)
        txt.delete('1.0', END)
        txt.insert(END, text)
elif config['main']['fileextension'] == '.xls':
    try:
        df = pd.read_excel(defaultdocument)        
        if (len(df)== 0):
            msg.showinfo('No records', 'No records')
        else:
            pass
        f2 = Frame(nfl6, height=200, width=300) 
        f2.pack(fill=BOTH,expand=1)
        table = Table(f2, dataframe=df,read_only=True)
        table.show()
        config.set('main','defaultdocument', qq)
    except FileNotFoundError as e:
        print(e)
        showerror('Error in opening file',e)
elif config['main']['fileextension'] == '.csv':
    try:
        df = pd.read_csv(defaultdocument)        
        if (len(df)== 0):
            msg.showinfo('No records', 'No records')
        else:
            pass
        f2 = Frame(nfl6, height=200, width=300) 
        f2.pack(fill=BOTH,expand=1)
        table = Table(f2, dataframe=df,read_only=True)
        table.show()
        config.set('main','defaultdocument', qq)
    except FileNotFoundError as e:
        print(e)
        showerror('Error in opening file',e)
elif config['main']['fileextension'] == ".docx":
    ee = docx2txt.process(qq)
    txt.delete('1.0', END)
    txt.insert(END, ee)
    config.set('main','defaultdocument', qq)
with open(loglocal, 'a') as j:
    j.write("Loading Config... \n")
smile = tkFont.Font(family=config['font']['font'], size=config['font']['size'], weight=config['font']['thickness'], slant=config['font']['incline'])
txt.config(font = smile)
txt.config(bg = config['color']['bg'])
txt.config(fg = config['color']['fg'])
with open(loglocal, 'a') as j:
    j.write("Loading Shell... \n")
root.mainloop()